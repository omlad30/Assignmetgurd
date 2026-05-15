from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()
    
    # Title Page
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('\n\n\nPROJECT REPORT\nON\n')
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    proj_name = document.add_paragraph()
    proj_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn_run = proj_name.add_run('AssignGuard: AI-Powered Secure Assignment Management System\n')
    pn_run.font.size = Pt(28)
    pn_run.bold = True
    
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run('\nSubmitted in partial fulfillment of the requirements for the degree of\nBachelor of Technology\n\n').font.size = Pt(14)
    
    document.add_page_break()
    
    # Abstract
    document.add_heading('ABSTRACT', level=1)
    document.add_paragraph(
        "AssignGuard is a comprehensive, modern Educational Technology (EdTech) platform designed to streamline the "
        "assignment submission process while maintaining rigorous academic integrity. It addresses the growing concern of "
        "plagiarism and AI-generated content in academic submissions by integrating advanced text extraction (OCR), "
        "similarity checking algorithms, and a Socratic AI Tutor. The platform serves three primary roles: Students, "
        "Teachers, and Administrators, providing a secure, responsive, and intuitive interface for classroom management, "
        "assignment tracking, and performance analytics."
    )
    document.add_page_break()

    # Table of Contents (Placeholder)
    document.add_heading('TABLE OF CONTENTS', level=1)
    document.add_paragraph("1. Introduction\n2. Problem Statement\n3. Proposed System\n4. System Architecture & Workflow\n5. Technologies Used\n6. Modules Description\n7. Screenshots\n8. Future Scope\n9. Conclusion")
    document.add_page_break()

    # 1. Introduction
    document.add_heading('1. INTRODUCTION', level=1)
    document.add_paragraph(
        "With the rapid digitization of education, managing assignments online has become the standard. "
        "However, this shift has introduced new challenges, primarily surrounding academic dishonesty. "
        "Students have unprecedented access to online resources and AI tools, making it difficult for educators "
        "to verify the originality of submissions. AssignGuard was developed to bridge this gap by providing a "
        "seamless digital classroom environment coupled with robust, automated security checks that prevent "
        "plagiarism bypasses and ensure fair evaluation."
    )

    # 2. Problem Statement
    document.add_heading('2. PROBLEM STATEMENT', level=1)
    document.add_paragraph(
        "Traditional Learning Management Systems (LMS) often lack built-in, sophisticated plagiarism detection "
        "that can handle diverse file formats (like image-based PDFs). Furthermore, educators spend an excessive "
        "amount of time manually grading and verifying the authenticity of assignments. There is a critical need for a "
        "platform that not only organizes classrooms and submissions but actively analyzes them for similarities and "
        "provides AI-driven constructive feedback to students before final submission."
    )

    # 3. Proposed System
    document.add_heading('3. PROPOSED SYSTEM', level=1)
    document.add_paragraph(
        "AssignGuard proposes a multi-tiered architecture that automates the verification of assignments. "
        "When a student submits a document, the system extracts the text (even from images using OCR) and compares "
        "it against a database of previous submissions using similarity algorithms. The system also features a "
        "Socratic AI Tutor that guides students rather than giving them direct answers, promoting genuine learning. "
        "Teachers receive a comprehensive dashboard highlighting 'Struggle Metrics' and plagiarism network graphs."
    )

    # 4. Technologies Used
    document.add_heading('4. TECHNOLOGIES USED', level=1)
    document.add_paragraph("Frontend Development:", style='List Bullet')
    document.add_paragraph("React.js (Vite): For a fast, component-based user interface.", style='List Bullet 2')
    document.add_paragraph("Tailwind CSS & Glassmorphism: For modern, responsive styling and animations.", style='List Bullet 2')
    document.add_paragraph("Lucide React: For consistent, scalable iconography.", style='List Bullet 2')
    
    document.add_paragraph("Backend Development:", style='List Bullet')
    document.add_paragraph("Node.js & Express.js: For building a scalable RESTful API.", style='List Bullet 2')
    document.add_paragraph("MongoDB (Atlas): NoSQL database for flexible data storage.", style='List Bullet 2')
    document.add_paragraph("JWT & bcryptjs: For secure authentication and password hashing.", style='List Bullet 2')
    
    document.add_paragraph("Security & Utility Services:", style='List Bullet')
    document.add_paragraph("Tesseract.js (OCR): To extract text from image-based PDFs, preventing hidden text bypasses.", style='List Bullet 2')
    document.add_paragraph("Google Gemini API: Powering the Socratic Tutor and AI-driven Pre-Flight checks.", style='List Bullet 2')
    document.add_paragraph("Socket.io: For real-time notifications of submissions and grading.", style='List Bullet 2')
    document.add_paragraph("Nodemailer: For secure OTP (One-Time Password) delivery.", style='List Bullet 2')

    # 5. System Architecture
    document.add_heading('5. SYSTEM ARCHITECTURE & WORKFLOW', level=1)
    document.add_paragraph(
        "The architecture follows a standard MERN stack paradigm with an integrated AI verification layer. "
        "Below is the textual representation of the workflow:"
    )
    workflow = (
        "1. User Authentication: Users log in via Email/Password or Google OAuth. Admins require 2FA OTP.\n"
        "2. Classroom Management: Teachers create classes and share invite codes. Students join using codes.\n"
        "3. Submission Flow: Student uploads a PDF -> File is parsed -> Text is extracted via OCR -> "
        "Pre-flight AI check runs -> Document is checked for similarity against DB -> Final Score generated.\n"
        "4. Real-time Feedback: Teacher receives Socket.io notification. Dashboard updates with new metrics."
    )
    document.add_paragraph(workflow)
    document.add_paragraph("[ PLACE WORKFLOW DIAGRAM IMAGE HERE ]", style='Intense Quote')

    # 6. Modules
    document.add_heading('6. MODULES DESCRIPTION', level=1)
    
    document.add_heading('A. Student Module', level=2)
    document.add_paragraph("Enables students to view enrolled classes, access pending assignments, utilize the Socratic Tutor for help, run pre-flight checks on drafts, and submit final PDFs.")
    
    document.add_heading('B. Teacher Module', level=2)
    document.add_paragraph("Allows educators to create classrooms, generate invite codes, post assignments, view real-time similarity scores, and analyze classroom 'Struggle Metrics'.")
    
    document.add_heading('C. Administrator Module', level=2)
    document.add_paragraph("Secured via Two-Factor Authentication (OTP), the admin dashboard oversees all platform activity, manages user roles, and monitors system health logs.")
    
    document.add_heading('D. Plagiarism & OCR Engine', level=2)
    document.add_paragraph("The core security engine that strips text from uploads, runs Tesseract OCR to prevent image-based cheating, and compares text vectors to flag copied content.")

    # 7. Screenshots
    document.add_page_break()
    document.add_heading('7. SCREENSHOTS', level=1)
    document.add_paragraph("1. Login / Registration Screen (Showcasing Glassmorphism Design)")
    document.add_paragraph("[ PLACE SCREENSHOT 1 HERE ]", style='Intense Quote')
    
    document.add_paragraph("2. Student Dashboard (Showing active assignments and progress)")
    document.add_paragraph("[ PLACE SCREENSHOT 2 HERE ]", style='Intense Quote')
    
    document.add_paragraph("3. Teacher Plagiarism Network Graph (Showing similarity links between students)")
    document.add_paragraph("[ PLACE SCREENSHOT 3 HERE ]", style='Intense Quote')
    
    document.add_paragraph("4. Admin Dashboard (Showing platform statistics)")
    document.add_paragraph("[ PLACE SCREENSHOT 4 HERE ]", style='Intense Quote')

    # 8. Future Scope
    document.add_page_break()
    document.add_heading('8. FUTURE SCOPE', level=1)
    document.add_paragraph("1. Multi-Language Support: Integrating translation APIs to detect plagiarism across different languages.", style='List Number')
    document.add_paragraph("2. Voice-Activated Socratic Tutor: Upgrading the AI tutor to accept voice inputs for accessibility.", style='List Number')
    document.add_paragraph("3. Automated Grading: Using AI to not only check for plagiarism but also grade assignments based on a teacher-provided rubric.", style='List Number')
    document.add_paragraph("4. Mobile Application: Developing a dedicated React Native application for iOS and Android.", style='List Number')

    # 9. Conclusion
    document.add_heading('9. CONCLUSION', level=1)
    document.add_paragraph(
        "AssignGuard successfully demonstrates how modern web technologies can be combined with Artificial Intelligence "
        "to solve real-world problems in education. By enforcing academic integrity through robust OCR and similarity "
        "checking, while simultaneously supporting students through the Socratic Tutor, the platform strikes a balance "
        "between security and learning. The architecture is scalable, user-friendly, and ready for deployment in "
        "real-world educational environments."
    )

    document.save('AssignGuard_Project_Report.docx')
    print("Report generated successfully as 'AssignGuard_Project_Report.docx'")

if __name__ == '__main__':
    create_report()
